"""
services/file_handler.py
Handles parsing of CSV, XLSX, and DOCX files into pandas DataFrames.
"""

import pandas as pd
import io
from docx import Document


def parse_file(filepath: str) -> pd.DataFrame:
    """
    Read a file from disk and return a pandas DataFrame.
    Supports: .csv, .xlsx, .xls, .docx
    """
    ext = filepath.rsplit(".", 1)[-1].lower()

    if ext == "csv":
        return _parse_csv(filepath)
    elif ext in ("xlsx", "xls"):
        return _parse_excel(filepath)
    elif ext == "docx":
        return _parse_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: .{ext}")


def _parse_csv(filepath: str) -> pd.DataFrame:
    """Parse CSV using pandas with smart encoding detection."""
    try:
        df = pd.read_csv(filepath, encoding="utf-8")
    except UnicodeDecodeError:
        df = pd.read_csv(filepath, encoding="latin-1")
    return df


def _parse_excel(filepath: str) -> pd.DataFrame:
    """Parse Excel file using openpyxl engine (supports .xlsx)."""
    return pd.read_excel(filepath, engine="openpyxl")


def _parse_docx(filepath: str) -> pd.DataFrame:
    """
    Extract the FIRST table found in a DOCX document and convert to DataFrame.
    The first row is used as the header.
    """
    doc = Document(filepath)
    if not doc.tables:
        raise ValueError("No tables found in the DOCX file.")

    table = doc.tables[0]
    data = []
    for row in table.rows:
        data.append([cell.text.strip() for cell in row.cells])

    if not data:
        raise ValueError("Table in DOCX is empty.")

    # First row as header
    headers = data[0]
    rows = data[1:]
    df = pd.DataFrame(rows, columns=headers)

    # Attempt to convert numeric columns automatically
    for col in df.columns:
        try:
            df[col] = pd.to_numeric(df[col])
        except (ValueError, TypeError):
            pass

    return df
